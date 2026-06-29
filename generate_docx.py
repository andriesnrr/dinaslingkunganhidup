import docx
import sys
import io

def generate_word_document(md_path, docx_path):
    doc = docx.Document()
    
    with io.open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            doc.add_heading(line[4:].replace('**', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('**', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].replace('**', ''), level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. '):
            doc.add_paragraph(line, style='List Number')
        elif line.startswith('---'):
            continue
        else:
            # Strip simple markdown bold/italics
            text = line.replace('**', '').replace('*', '')
            doc.add_paragraph(text)
            
    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

md_file = "C:/Users/andri/.gemini/antigravity-ide/brain/9a7e2bf3-abf6-45f5-b918-42aabbae8034/artifacts/draf_laporan_magang.md"
docx_file = "Draf_Laporan_Magang_DLH.docx"

generate_word_document(md_file, docx_file)
